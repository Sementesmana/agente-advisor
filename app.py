# -*- coding: utf-8 -*-
"""agente-advisor — Advisor Alfredo Soares (Sementes Maná)
Pipeline: coletar transcrição YouTube -> sintetizar (Claude via mana-llm-gateway) -> consolidar mente -> chat persona.
Estado derivado do filesystem (data/): transcricoes/{id}.txt, sinteses/{id}.md, mente/{tema}.md, consolidado.json
"""
import os, io, json, re, glob, threading, traceback, datetime
import requests
from flask import Flask, request, jsonify, Response

BASE = os.path.dirname(os.path.abspath(__file__))
DATA = os.environ.get('DATA_DIR', os.path.join(BASE, 'data'))
_SEED = os.path.join(BASE, 'data')
# Volume persistente (Railway): se DATA_DIR aponta pra um volume vazio, semeia com o data/ do repo
if DATA != _SEED and not os.path.exists(os.path.join(DATA, 'videos.json')):
    import shutil
    os.makedirs(DATA, exist_ok=True)
    shutil.copytree(_SEED, DATA, dirs_exist_ok=True)
    print('[boot] volume semeado a partir do repo', flush=True)
GW_URL = os.environ.get('LLM_GATEWAY_URL', '').rstrip('/')
GW_KEY = os.environ.get('LLM_GATEWAY_KEY', '')
MODEL = os.environ.get('LLM_MODEL', 'claude-sonnet-4-5')                 # chat do Advisor (qualidade)
MODEL_SINTESE = os.environ.get('LLM_MODEL_SINTESE', 'mana-rapido')       # síntese/consolidação (barato = Haiku)
CRON_HORA = int(os.environ.get('CRON_HORA', '7'))  # BRT
YT_CHANNEL_ID = os.environ.get('YT_CHANNEL_ID', 'UCh9HMS4C3F02msM-kiilAdA')  # @canaldoalfredosoares
PROXY_URL = os.environ.get('PROXY_URL', '')  # proxy residencial p/ YouTube (http://user:pass@host:porta)
PROXIES = {'http': PROXY_URL, 'https': PROXY_URL} if PROXY_URL else None
TAXONOMIA = ['modelo-de-negocio','vendas-e-ofertas','marketing-de-influencia','canais-e-varejo',
             'conteudo-e-audiencia','gestao-e-pessoas','mentalidade-empreendedora',
             'networking-e-conexoes','branding-e-posicionamento']

app = Flask(__name__)
PROGRESSO = {'rodando': False, 'log': [], 'abortar': False}

@app.before_request
def _preflight():
    if request.method == 'OPTIONS' and request.path.startswith('/api/'):
        r = Response('')
        r.headers['Access-Control-Allow-Origin'] = '*'
        r.headers['Access-Control-Allow-Methods'] = 'GET, POST, OPTIONS'
        r.headers['Access-Control-Allow-Headers'] = 'content-type'
        return r

@app.after_request
def _cors(resp):
    if request.path.startswith('/api/'):
        resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp

def p(*a): return os.path.join(DATA, *a)
def ler(path): return open(path, encoding='utf-8').read() if os.path.exists(path) else ''
def gravar(path, txt):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    open(path, 'w', encoding='utf-8').write(txt)

def videos():
    vs = json.loads(ler(p('videos.json')) or '[]')
    ign = set(json.loads(ler(p('ignorados.json')) or '[]'))
    vs = [v for v in vs if v['id'] not in ign]
    cons = set(json.loads(ler(p('consolidado.json')) or '[]'))
    for v in vs:
        if v['id'] in cons: v['status'] = 'consolidado'
        elif os.path.exists(p('sinteses', v['id'] + '.md')): v['status'] = 'sintetizado'
        elif os.path.exists(p('transcricoes', v['id'] + '.txt')): v['status'] = 'transcrito'
        else: v['status'] = 'pendente'
    return vs

def log(msg):
    PROGRESSO['log'].append(msg)
    PROGRESSO['log'] = PROGRESSO['log'][-80:]
    print('[pipeline]', msg, flush=True)

# ---------- LLM (mana-llm-gateway, API compatível Anthropic /v1/messages) ----------
def _gw_headers():
    return {'x-api-key': GW_KEY, 'Authorization': 'Bearer ' + GW_KEY,
            'anthropic-version': '2023-06-01', 'content-type': 'application/json'}

def llm(system, user, max_tokens=8000, model=None):
    r = requests.post(GW_URL + '/v1/messages',
        headers=_gw_headers(),
        json={'model': model or MODEL_SINTESE, 'max_tokens': max_tokens, 'system': system,
              'messages': [{'role': 'user', 'content': user}]}, timeout=300)
    r.raise_for_status()
    return ''.join(b.get('text', '') for b in r.json().get('content', []))

# ---------- 1 COLETOR (determinístico) ----------
import html as _html

def ignorar(vid, motivo):
    ign = json.loads(ler(p('ignorados.json')) or '[]')
    if vid not in ign:
        ign.append(vid); gravar(p('ignorados.json'), json.dumps(ign))
    log('Ignorado %s (%s)' % (vid, motivo))

def dur_seg(vid):
    """Duração do vídeo em segundos (0 se não conseguir determinar → deixa passar)."""
    try:
        s = requests.Session()
        if PROXIES: s.proxies = PROXIES
        s.headers['User-Agent'] = 'com.google.android.youtube/20.10.38 (Linux; U; Android 11) gzip'
        j = s.post('https://www.youtube.com/youtubei/v1/player', json={
            'context': {'client': {'clientName': 'ANDROID', 'clientVersion': '20.10.38',
                                   'androidSdkVersion': 30, 'hl': 'pt', 'gl': 'BR'}},
            'videoId': vid}, timeout=30).json()
        return int(j.get('videoDetails', {}).get('lengthSeconds', 0) or 0)
    except Exception:
        return 0

def coletar():
    """Só RSS — rápido, traz todos os vídeos do canal pra fila (a limpeza de shorts é botão separado)."""
    r = requests.get('https://www.youtube.com/feeds/videos.xml?channel_id=' + YT_CHANNEL_ID,
                     timeout=30, headers={'User-Agent': 'Mozilla/5.0'}, proxies=PROXIES)
    r.raise_for_status()
    entradas = re.findall(r'<entry>([\s\S]*?)</entry>', r.text)
    vs = json.loads(ler(p('videos.json')) or '[]')
    conhecidos = {v['id'] for v in vs}
    novos, feed_ids = [], []
    for e in entradas:  # feed vem do mais recente pro mais antigo
        vid = (re.search(r'<yt:videoId>([\w-]+)</yt:videoId>', e) or [None, None])[1]
        tit = (re.search(r'<title>([\s\S]*?)</title>', e) or [None, ''])[1]
        pub = (re.search(r'<published>([\d-]+)', e) or [None, ''])[1]
        if not vid: continue
        feed_ids.append(vid)
        if vid not in conhecidos:
            item = {'id': vid, 'titulo': _html.unescape(tit).strip(), 'views': '', 'data': pub}
            d = dur_seg(vid)              # traz a duração já no Buscar (só p/ os novos = poucos)
            if d > 0: item['dur'] = d
            novos.append(item)
        else:
            for v in vs:
                if v['id'] == vid and not v.get('data'): v['data'] = pub
    vs = novos + vs
    pos = {vid: i for i, vid in enumerate(feed_ids)}
    vs.sort(key=lambda v: pos.get(v['id'], 10**6))
    gravar(p('videos.json'), json.dumps(vs, ensure_ascii=False, indent=1))
    log('Coletor: %d vídeo(s) novo(s) no canal' % len(novos))
    return len(novos)

def limpar_shorts():
    """REGRA LOCAL pura — remove da fila quem já tem duração < 4min. Não fala com o YouTube, é instantâneo."""
    PROGRESSO['log'] = []
    vs = json.loads(ler(p('videos.json')) or '[]')
    rem = [v['id'] for v in vs if 0 < (v.get('dur') or 0) < 240]
    for vid in rem: ignorar(vid, 'short (regra local)')
    vs2 = [v for v in vs if v['id'] not in set(rem)]
    gravar(p('videos.json'), json.dumps(vs2, ensure_ascii=False, indent=1))
    sem = sum(1 for v in vs2 if not v.get('dur'))
    log('🧹 %d short(s) removido(s) pela duração.' % len(rem) +
        ((' %d vídeo(s) ainda sem duração — clique em Buscar (traz a duração dos novos).' % sem) if sem else ''))

def transcrever(vid):
    """3 tentativas com pausa — o YouTube bloqueia IP de datacenter de forma intermitente."""
    import time
    ult = None
    for _ in range(3):
        try:
            return _transcrever_1x(vid)
        except RuntimeError as e:
            if str(e) == 'CURTO': raise
            ult = e; time.sleep(10)
    raise ult

def _transcrever_1x(vid):
    """Legenda automática via innertube (client ANDROID). Levanta exceção se bloqueado/sem legenda."""
    s = requests.Session()
    if PROXIES: s.proxies = PROXIES
    s.headers['User-Agent'] = 'com.google.android.youtube/20.10.38 (Linux; U; Android 11) gzip'
    j = s.post('https://www.youtube.com/youtubei/v1/player', json={
        'context': {'client': {'clientName': 'ANDROID', 'clientVersion': '20.10.38',
                               'androidSdkVersion': 30, 'hl': 'pt', 'gl': 'BR'}},
        'videoId': vid}, timeout=60).json()
    tracks = (j.get('captions', {}).get('playerCaptionsTracklistRenderer', {}).get('captionTracks', []))
    tr = next((t for t in tracks if t.get('languageCode', '').startswith('pt')), tracks[0] if tracks else None)
    if not tr: raise RuntimeError('SEM_LEGENDA')
    dur = int(j.get('videoDetails', {}).get('lengthSeconds', 0) or 0)
    if 0 < dur < 240: raise RuntimeError('CURTO')  # shorts/teasers não entram no cérebro
    xml = s.get(tr['baseUrl'], timeout=60).text
    dec = lambda x: x.replace('&amp;','&').replace('&lt;','<').replace('&gt;','>').replace('&#39;',"'").replace('&quot;','"')
    parts = [dec(re.sub(r'<[^>]+>', ' ', m.group(1))) for m in re.finditer(r'<(?:text|p)[^>]*>([\s\S]*?)</(?:text|p)>', xml)]
    txt = re.sub(r'\s+', ' ', ' '.join(parts)).strip()
    if len(txt) < 2500: raise RuntimeError('CURTO')
    return txt

# ---------- 2 ANALISTA (probabilístico) ----------
SINTESE_SYS = """Você sintetiza vídeos do Alfredo Soares (co-fundador do G4 Educação) para uma base de conhecimento de negócios.
Responda APENAS com o markdown da síntese, em português, neste formato exato:

# {titulo}
url: https://www.youtube.com/watch?v={id}
views: {views}

## Resumo
(1-2 parágrafos densos)

## Conceitos e frameworks
- **Nome** — explicação (marque de quem é a lição se houver convidado)

## Insights acionáveis
- (5-12 itens práticos)

## Cases e números citados
- (empresas, pessoas, métricas)

## Frases marcantes
- "máx. 3 citações, cada uma com menos de 15 palavras"

## Temas
tags: (3-6 tags, OBRIGATORIAMENTE escolhidas desta lista: %s)""" % ', '.join(TAXONOMIA)

def sintetizar(v):
    txt = ler(p('transcricoes', v['id'] + '.txt'))
    md = llm(SINTESE_SYS, 'Vídeo: %s (id %s, %s views)\n\nTRANSCRIÇÃO:\n%s' % (v['titulo'], v['id'], v.get('views',''), txt[:180000]))
    gravar(p('sinteses', v['id'] + '.md'), md.strip())
    m = re.search(r'tags:\s*(.+)', md)
    return [t.strip() for t in m.group(1).split(',') if t.strip() in TAXONOMIA] if m else []

# ---------- 3 CONSOLIDADOR (probabilístico, incremental) ----------
CONSOL_SYS = """Você mantém a "mente do Alfredo Soares": arquivos de tópico com princípios numerados.
Receberá o arquivo atual do tópico e novas sínteses de vídeos. Reescreva o ARQUIVO COMPLETO do tópico:
- Mantenha o formato: título, princípios como **N. Título.** corpo com citação [Nome do vídeo · VIDEOID] e seção ## Fontes ao final.
- Funda repetições (o que se repete em vários vídeos ganha destaque como princípio central); divergências viram nuance; o vídeo mais recente prevalece em conflito.
- Não invente nada que não esteja nas fontes. Responda APENAS com o markdown do arquivo."""

def consolidar(tema, novos_ids):
    atual = ler(p('mente', tema + '.md'))
    sints = '\n\n---\n\n'.join(ler(p('sinteses', i + '.md')) for i in novos_ids)
    md = llm(CONSOL_SYS, 'TÓPICO: %s\n\nARQUIVO ATUAL:\n%s\n\nNOVAS SÍNTESES:\n%s' % (tema, atual or '(novo tópico)', sints), 12000)
    gravar(p('mente', tema + '.md'), md.strip())

# ---------- PIPELINE ----------
def processar(ids=None):
    if PROGRESSO['rodando']: return
    PROGRESSO['rodando'] = True; PROGRESSO['log'] = []; PROGRESSO['abortar'] = False
    try:
        if not ids:
            try: coletar()
            except Exception as e: log('Coletor falhou (segue com a fila atual): %s' % e)
        temas_novos = {}  # tema -> [video ids]
        for v in videos():
            if PROGRESSO['abortar']:
                log('⛔ Abortado pelo usuário — o que já foi sintetizado consolida na próxima rodada.'); break
            if ids and v['id'] not in ids: continue
            try:
                if v['status'] == 'pendente':
                    log('Transcrevendo: ' + v['titulo'][:60])
                    try:
                        txt = transcrever(v['id'])
                    except RuntimeError as e:
                        if str(e) == 'CURTO': ignorar(v['id'], 'curto/short'); continue
                        raise
                    gravar(p('transcricoes', v['id'] + '.txt'), v['titulo'] + '\nhttps://www.youtube.com/watch?v=' + v['id'] + '\n' + txt)
                    v['status'] = 'transcrito'
                if v['status'] == 'transcrito':
                    log('Sintetizando: ' + v['titulo'][:60])
                    for t in sintetizar(v): temas_novos.setdefault(t, []).append(v['id'])
                    v['status'] = 'sintetizado'
                elif v['status'] == 'sintetizado':  # sintetizado mas nunca consolidado
                    md = ler(p('sinteses', v['id'] + '.md'))
                    m = re.search(r'tags:\s*(.+)', md)
                    for t in ([x.strip() for x in m.group(1).split(',')] if m else []):
                        if t in TAXONOMIA: temas_novos.setdefault(t, []).append(v['id'])
            except Exception as e:
                log('ERRO %s: %s' % (v['id'], e))
        if PROGRESSO['abortar']: temas_novos = {}
        for tema, ids in temas_novos.items():
            if PROGRESSO['abortar']: break
            log('Consolidando tópico: ' + tema)
            try: consolidar(tema, ids)
            except Exception as e: log('ERRO consolidação %s: %s' % (tema, e)); continue
        cons = set(json.loads(ler(p('consolidado.json')) or '[]'))
        cons |= {i for ids in temas_novos.values() for i in ids}
        gravar(p('consolidado.json'), json.dumps(sorted(cons)))
        log('Pipeline concluído.')
    except Exception:
        log('ERRO geral: ' + traceback.format_exc()[-300:])
    finally:
        PROGRESSO['rodando'] = False

# ---------- CONTEXTO DA EMPRESA ----------
def extrair_texto(nome, dados):
    n = nome.lower()
    if n.endswith('.pdf'):
        from pypdf import PdfReader
        return '\n'.join(pg.extract_text() or '' for pg in PdfReader(io.BytesIO(dados)).pages)
    if n.endswith('.docx'):
        import docx
        d = docx.Document(io.BytesIO(dados))
        partes = [par.text for par in d.paragraphs]
        for t in d.tables:
            for row in t.rows: partes.append(' | '.join(c.text for c in row.cells))
        return '\n'.join(partes)
    if n.endswith(('.html', '.htm')):
        return re.sub(r'\s+', ' ', re.sub(r'<script[\s\S]*?</script>|<style[\s\S]*?</style>|<[^>]+>', ' ',
                                          dados.decode('utf-8', 'ignore')))
    return dados.decode('utf-8', 'ignore')

# ---------- CONTEXTO: Empresa -> Área -> itens ----------
def _slug(s):
    s = re.sub(r'[^\w\s-]', '', (s or '').strip().lower())
    return re.sub(r'[\s_-]+', '-', s)[:50] or 'sem-nome'

CTX = 'ctx'  # data/ctx/

def empresas():
    return json.loads(ler(p(CTX, 'empresas.json')) or '[]')

def salvar_empresas(lst):
    gravar(p(CTX, 'empresas.json'), json.dumps(lst, ensure_ascii=False, indent=1))

def garantir_empresa_padrao():
    if not empresas():
        salvar_empresas([{'slug': 'sementes-mana', 'nome': 'Sementes Maná', 'ativa': True}])
        gravar(p(CTX, 'sementes-mana', 'perfil.md'), '')

def empresa_ativa():
    es = empresas()
    return next((e for e in es if e.get('ativa')), es[0] if es else None)

def ctx_arvore(slug):
    """Retorna {perfil, areas:[{slug,nome,nota,docs:[{nome,chars}]}]} de uma empresa."""
    base = p(CTX, slug)
    perfil = ler(os.path.join(base, 'perfil.md'))
    areas = []
    if os.path.isdir(base):
        for a in sorted(os.listdir(base)):
            ad = os.path.join(base, a)
            if not os.path.isdir(ad): continue
            docs = [{'nome': os.path.basename(f)[:-4], 'chars': len(ler(f))}
                    for f in sorted(glob.glob(os.path.join(ad, '*.txt')))
                    if not os.path.basename(f).startswith('_')]
            nome = ler(os.path.join(ad, '_nome.txt')) or a
            areas.append({'slug': a, 'nome': nome.strip() or a, 'nota': ler(os.path.join(ad, '_nota.md')), 'docs': docs})
    return {'perfil': perfil, 'areas': areas}

def _empresa_por_slug(slug):
    es = empresas()
    return next((x for x in es if x['slug'] == slug), None) or empresa_ativa()

def contexto_para_chat(empresa_slug=None, area_slug=None):
    """Monta o contexto de uma empresa (default: ativa). Se area_slug, foca só naquela área."""
    e = _empresa_por_slug(empresa_slug)
    if not e: return ''
    arv = ctx_arvore(e['slug'])
    partes = ['EMPRESA: ' + e['nome']]
    if arv['perfil'].strip(): partes.append('PERFIL GERAL:\n' + arv['perfil'].strip())
    for a in arv['areas']:
        if area_slug and a['slug'] != area_slug: continue
        bloco = ['--- ÁREA: %s ---' % a['nome']]
        if a['nota'].strip(): bloco.append(a['nota'].strip())
        for f in sorted(glob.glob(p(CTX, e['slug'], a['slug'], '*.txt'))):
            if os.path.basename(f).startswith('_'): continue
            bloco.append('[doc: %s]\n%s' % (os.path.basename(f)[:-4], ler(f)[:25000]))
        if len(bloco) > 1: partes.append('\n'.join(bloco))
    return ('\n\n'.join(partes))[:150000]

# ---------- 4 PERSONA / CHAT ----------
def chat(pergunta, historico, empresa_slug=None, area_slug=None):
    persona = ler(p('mente', 'persona.md'))
    mente = '\n\n'.join(ler(f) for f in sorted(glob.glob(p('mente', '*.md'))) if not f.endswith('persona.md'))
    ctx = contexto_para_chat(empresa_slug, area_slug)
    sys = persona + '\n\n=== BASE DE CONHECIMENTO (a mente) ===\n' + mente + \
          (('\n\n=== CONTEXTO DA EMPRESA (use e cite a área de origem quando relevante) ===\n' + ctx) if ctx else '') + \
          '\n\nResponda como o advisor: direto, provocador, com plano de ação e a conta feita. Use o contexto da empresa quando existir — a orientação deve ser específica pro negócio. Cite os vídeos-fonte (nome + link) dos princípios que usar.'
    msgs = historico[-8:] + [{'role': 'user', 'content': pergunta}]
    r = requests.post(GW_URL + '/v1/messages',
        headers=_gw_headers(),
        json={'model': MODEL, 'max_tokens': 3000, 'system': sys, 'messages': msgs}, timeout=180)
    r.raise_for_status()
    return ''.join(b.get('text', '') for b in r.json().get('content', []))

# ---------- API ----------
@app.route('/api/estado')
def api_estado():
    topics = []
    for f in sorted(glob.glob(p('mente', '*.md'))):
        slug = os.path.basename(f)[:-3]
        if slug == 'persona': continue
        raw = ler(f)
        topics.append({'slug': slug, 'title': (re.search(r'^# (.+)$', raw, re.M) or [None,'?'])[1] if re.search(r'^# (.+)$', raw, re.M) else slug,
                       'n': len(re.findall(r'\*\*\d+\.', raw))})
    return jsonify({'videos': videos(), 'topics': topics, 'progresso': PROGRESSO,
                    'gateway': bool(GW_URL and GW_KEY)})

@app.route('/api/mente/<slug>')
def api_mente(slug):
    if not re.match(r'^[\w-]+$', slug): return ('', 404)
    return jsonify({'md': ler(p('mente', slug + '.md'))})

@app.route('/api/transcricao', methods=['POST', 'OPTIONS'])
def api_transcricao():
    """Plano B: recebe transcrição extraída no navegador do usuário (quando o YouTube bloqueia o IP do servidor)."""
    resp_headers = {'Access-Control-Allow-Origin': 'https://www.youtube.com',
                    'Access-Control-Allow-Methods': 'POST',
                    'Access-Control-Allow-Headers': 'content-type'}
    if request.method == 'OPTIONS':
        return Response('', headers=resp_headers)
    d = request.get_json(force=True)
    vid, txt = d.get('id', ''), (d.get('texto', '') or '').strip()
    if not re.match(r'^[\w-]{11}$', vid): return jsonify({'erro': 'id inválido'}), 400
    if len(txt) < 2500:
        ignorar(vid, 'curto/short via navegador')
        return jsonify({'ok': True, 'ignorado': True}, ), 200, resp_headers
    vs = {v['id']: v for v in json.loads(ler(p('videos.json')) or '[]')}
    tit = vs.get(vid, {}).get('titulo', vid)
    gravar(p('transcricoes', vid + '.txt'), tit + '\nhttps://www.youtube.com/watch?v=' + vid + '\n' + txt)
    return jsonify({'ok': True, 'chars': len(txt)}), 200, resp_headers

@app.route('/api/limpar-shorts', methods=['POST'])
def api_limpar_shorts():
    limpar_shorts()  # regra local, instantâneo
    return jsonify({'ok': True})

@app.route('/api/set-duracoes', methods=['POST'])
def api_set_duracoes():
    """Recebe {id: segundos} extraídos pelo navegador (IP residencial) e grava a duração na fila."""
    m = (request.get_json(force=True) or {}).get('duracoes', {})
    vs = json.loads(ler(p('videos.json')) or '[]')
    n = 0
    for v in vs:
        if v['id'] in m:
            v['dur'] = int(m[v['id']]); n += 1
    gravar(p('videos.json'), json.dumps(vs, ensure_ascii=False, indent=1))
    return jsonify({'ok': True, 'atualizados': n})

@app.route('/api/set-meta', methods=['POST'])
def api_set_meta():
    """Corrige título/duração/fonte de vídeos (ex.: os que ficaram sem nome ao adicionar sob bloqueio)."""
    m = (request.get_json(force=True) or {}).get('metas', {})  # {id: {titulo, dur, fonte}}
    vs = json.loads(ler(p('videos.json')) or '[]')
    n = 0
    for v in vs:
        if v['id'] in m:
            mm = m[v['id']]
            if mm.get('titulo'): v['titulo'] = mm['titulo']
            if mm.get('dur'): v['dur'] = int(mm['dur'])
            if mm.get('fonte'): v['fonte'] = mm['fonte']
            n += 1
    gravar(p('videos.json'), json.dumps(vs, ensure_ascii=False, indent=1))
    return jsonify({'ok': True, 'atualizados': n})

@app.route('/api/sem-titulo')
def api_sem_titulo():
    """Lista vídeos cujo título ficou igual ao id (pro navegador re-buscar os nomes)."""
    vs = json.loads(ler(p('videos.json')) or '[]')
    return jsonify([v['id'] for v in vs if v.get('titulo', '') == v['id']])

# ---------- Síntese feita no Cowork (subagentes) — app só grava, sem LLM ----------
@app.route('/api/transcricao/<vid>')
def api_get_transcricao(vid):
    if not re.match(r'^[\w-]{11}$', vid): return ('', 404)
    txt = ler(p('transcricoes', vid + '.txt'))
    if not txt: return jsonify({}), 404
    tit = next((v['titulo'] for v in json.loads(ler(p('videos.json')) or '[]') if v['id'] == vid), vid)
    return jsonify({'id': vid, 'titulo': tit, 'texto': txt})

@app.route('/api/pendentes-sintese')
def api_pendentes_sintese():
    """IDs que têm transcrição mas ainda não têm síntese."""
    out = []
    for f in sorted(glob.glob(p('transcricoes', '*.txt'))):
        vid = os.path.basename(f)[:-4]
        if not os.path.exists(p('sinteses', vid + '.md')):
            tit = next((v['titulo'] for v in json.loads(ler(p('videos.json')) or '[]') if v['id'] == vid), vid)
            out.append({'id': vid, 'titulo': tit})
    return jsonify(out)

@app.route('/api/set-sintese/<vid>', methods=['POST'])
def api_set_sintese(vid):
    """Recebe a síntese .md já pronta (feita no Cowork) e grava — sem chamar LLM."""
    if not re.match(r'^[\w-]{11}$', vid): return ('', 404)
    gravar(p('sinteses', vid + '.md'), (request.get_json(force=True) or {}).get('md', '').strip())
    return jsonify({'ok': True})

@app.route('/api/importar-lote', methods=['POST'])
def api_importar_lote():
    """Recebe TUDO de uma vez (sínteses + mente + ids consolidados), feito no Cowork. Grava sem LLM."""
    d = request.get_json(force=True) or {}
    n_s = n_m = 0
    for vid, md in (d.get('sinteses', {}) or {}).items():
        if re.match(r'^[\w-]{11}$', vid): gravar(p('sinteses', vid + '.md'), (md or '').strip()); n_s += 1
    for tema, md in (d.get('mente', {}) or {}).items():
        if re.match(r'^[\w-]+$', tema): gravar(p('mente', tema + '.md'), (md or '').strip()); n_m += 1
    ids = d.get('consolidado', [])
    if ids:
        cons = set(json.loads(ler(p('consolidado.json')) or '[]')) | set(ids)
        gravar(p('consolidado.json'), json.dumps(sorted(cons)))
    return jsonify({'ok': True, 'sinteses': n_s, 'temas': n_m})

@app.route('/api/set-mente/<tema>', methods=['POST'])
def api_set_mente(tema):
    """Recebe o arquivo .md de um tópico da mente (consolidado no Cowork) e grava; marca ids como consolidados."""
    if not re.match(r'^[\w-]+$', tema): return ('', 404)
    d = request.get_json(force=True) or {}
    gravar(p('mente', tema + '.md'), (d.get('md', '') or '').strip())
    ids = d.get('ids', [])
    if ids:
        cons = set(json.loads(ler(p('consolidado.json')) or '[]')) | set(ids)
        gravar(p('consolidado.json'), json.dumps(sorted(cons)))
    return jsonify({'ok': True})

@app.route('/api/ignorar', methods=['POST'])
def api_ignorar():
    """Remove uma lista de IDs da fila e manda pros ignorados (usado p/ limpar shorts identificados fora)."""
    ids = set((request.get_json(force=True) or {}).get('ids', []))
    vs = json.loads(ler(p('videos.json')) or '[]')
    manter = [v for v in vs if v['id'] not in ids]
    gravar(p('videos.json'), json.dumps(manter, ensure_ascii=False, indent=1))
    for vid in ids: ignorar(vid, 'short (lote)')
    return jsonify({'ok': True, 'removidos': len(vs) - len(manter)})

@app.route('/api/abortar', methods=['POST'])
def api_abortar():
    PROGRESSO['abortar'] = True
    return jsonify({'ok': True})

def yt_meta(vid):
    import time
    for tent in range(3):
        try:
            s = requests.Session()
            if PROXIES: s.proxies = PROXIES
            s.headers['User-Agent'] = 'com.google.android.youtube/20.10.38 (Linux; U; Android 11) gzip'
            j = s.post('https://www.youtube.com/youtubei/v1/player', json={
                'context': {'client': {'clientName': 'ANDROID', 'clientVersion': '20.10.38',
                                       'androidSdkVersion': 30, 'hl': 'pt', 'gl': 'BR'}},
                'videoId': vid}, timeout=30).json()
            d = j.get('videoDetails', {})
            if d.get('title'):
                return {'titulo': d['title'], 'dur': int(d.get('lengthSeconds', 0) or 0), 'autor': d.get('author', '')}
        except Exception:
            pass
        if tent < 2: time.sleep(4)
    return {'titulo': vid, 'dur': 0, 'autor': ''}

def extrair_video_id(tok):
    tok = tok.strip()
    m = re.search(r'(?:v=|youtu\.be/|shorts/|embed/|/live/)([\w-]{11})', tok)
    if m: return m.group(1)
    return tok if re.match(r'^[\w-]{11}$', tok) else None

@app.route('/api/add-video', methods=['POST'])
def api_add_video():
    txt = (request.get_json(force=True) or {}).get('urls', '')
    ids, vistos = [], set()
    for tok in re.split(r'[\s,]+', txt.strip()):
        vid = extrair_video_id(tok)
        if vid and vid not in vistos: ids.append(vid); vistos.add(vid)
    if not ids: return jsonify({'erro': 'nenhuma URL válida'}), 400
    vs = json.loads(ler(p('videos.json')) or '[]')
    conhecidos = {v['id'] for v in vs}
    add, novos = 0, []
    for vid in ids:
        if vid in conhecidos: continue
        try: m = yt_meta(vid)
        except Exception: m = {'titulo': vid, 'dur': 0, 'autor': ''}
        item = {'id': vid, 'titulo': m['titulo'], 'views': '', 'data': '', 'fonte': m.get('autor') or 'avulso'}
        if m['dur'] > 0: item['dur'] = m['dur']
        novos.append(item); conhecidos.add(vid); add += 1
    vs = novos + vs
    gravar(p('videos.json'), json.dumps(vs, ensure_ascii=False, indent=1))
    return jsonify({'ok': True, 'adicionados': add, 'total_validas': len(ids)})

@app.route('/api/coletar', methods=['POST'])
def api_coletar():
    try:
        n = coletar()
        return jsonify({'ok': True, 'novos': n})
    except Exception as e:
        return jsonify({'erro': str(e)}), 500

@app.route('/api/processar', methods=['POST'])
def api_processar():
    ids = (request.get_json(silent=True) or {}).get('ids')
    threading.Thread(target=processar, args=(ids,), daemon=True).start()
    return jsonify({'ok': True})

@app.route('/api/ordem', methods=['POST'])
def api_ordem():
    ids = (request.get_json(force=True) or {}).get('ids', [])
    pos = {vid: i for i, vid in enumerate(ids)}
    vs = json.loads(ler(p('videos.json')) or '[]')
    vs.sort(key=lambda v: pos.get(v['id'], 10**6))  # sort estável: não listados mantêm ordem
    gravar(p('videos.json'), json.dumps(vs, ensure_ascii=False, indent=1))
    return jsonify({'ok': True})

@app.route('/api/contexto')
def api_contexto():
    garantir_empresa_padrao()
    e = empresa_ativa()
    return jsonify({'empresas': empresas(), 'ativa': e['slug'] if e else None,
                    'arvore': ctx_arvore(e['slug']) if e else {'perfil': '', 'areas': []}})

@app.route('/api/contexto/empresa', methods=['POST'])
def api_ctx_empresa_nova():
    nome = (request.get_json(force=True) or {}).get('nome', '').strip()
    if not nome: return jsonify({'erro': 'informe o nome'}), 400
    es = empresas(); slug = _slug(nome)
    if any(x['slug'] == slug for x in es): return jsonify({'erro': 'empresa já existe'}), 400
    for x in es: x['ativa'] = False
    es.append({'slug': slug, 'nome': nome, 'ativa': True})
    salvar_empresas(es); gravar(p(CTX, slug, 'perfil.md'), '')
    return jsonify({'ok': True, 'slug': slug})

@app.route('/api/contexto/empresa-ativa', methods=['POST'])
def api_ctx_empresa_ativa():
    slug = (request.get_json(force=True) or {}).get('slug', '')
    es = empresas()
    if not any(x['slug'] == slug for x in es): return jsonify({'erro': 'não achei'}), 404
    for x in es: x['ativa'] = (x['slug'] == slug)
    salvar_empresas(es); return jsonify({'ok': True})

@app.route('/api/contexto/empresa/<slug>', methods=['DELETE'])
def api_ctx_empresa_del(slug):
    if not re.match(r'^[\w-]+$', slug): return ('', 404)
    es = [x for x in empresas() if x['slug'] != slug]
    if es and not any(x.get('ativa') for x in es): es[0]['ativa'] = True
    salvar_empresas(es)
    import shutil
    if os.path.isdir(p(CTX, slug)): shutil.rmtree(p(CTX, slug))
    return jsonify({'ok': True})

@app.route('/api/contexto/perfil', methods=['POST'])
def api_ctx_perfil():
    d = request.get_json(force=True) or {}
    slug = d.get('empresa', '')
    if not re.match(r'^[\w-]+$', slug): return ('', 404)
    gravar(p(CTX, slug, 'perfil.md'), d.get('perfil', ''))
    return jsonify({'ok': True})

@app.route('/api/contexto/area', methods=['POST'])
def api_ctx_area_nova():
    d = request.get_json(force=True) or {}
    emp, nome = d.get('empresa', ''), (d.get('nome', '') or '').strip()
    if not re.match(r'^[\w-]+$', emp) or not nome: return jsonify({'erro': 'dados'}), 400
    aslug = _slug(nome)
    gravar(p(CTX, emp, aslug, '_nome.txt'), nome)
    if not os.path.exists(p(CTX, emp, aslug, '_nota.md')): gravar(p(CTX, emp, aslug, '_nota.md'), '')
    return jsonify({'ok': True, 'slug': aslug})

@app.route('/api/contexto/area/<emp>/<area>', methods=['DELETE'])
def api_ctx_area_del(emp, area):
    if not re.match(r'^[\w-]+$', emp) or not re.match(r'^[\w-]+$', area): return ('', 404)
    import shutil
    if os.path.isdir(p(CTX, emp, area)): shutil.rmtree(p(CTX, emp, area))
    return jsonify({'ok': True})

@app.route('/api/contexto/nota', methods=['POST'])
def api_ctx_nota():
    d = request.get_json(force=True) or {}
    emp, area = d.get('empresa', ''), d.get('area', '')
    if not re.match(r'^[\w-]+$', emp) or not re.match(r'^[\w-]+$', area): return ('', 404)
    gravar(p(CTX, emp, area, '_nota.md'), d.get('nota', ''))
    return jsonify({'ok': True})

@app.route('/api/contexto/upload', methods=['POST'])
def api_ctx_upload():
    emp, area = request.form.get('empresa', ''), request.form.get('area', '')
    if not re.match(r'^[\w-]+$', emp) or not re.match(r'^[\w-]+$', area):
        return jsonify({'erro': 'escolha empresa e área'}), 400
    f = request.files.get('arquivo')
    if not f: return jsonify({'erro': 'sem arquivo'}), 400
    dados = f.read()
    if len(dados) > 15 * 1024 * 1024: return jsonify({'erro': 'máx 15MB'}), 400
    try: txt = extrair_texto(f.filename, dados)
    except Exception as e: return jsonify({'erro': 'falha ao extrair: %s' % e}), 400
    nome = re.sub(r'[^\w.-]+', '_', os.path.splitext(f.filename)[0])[:60]
    gravar(p(CTX, emp, area, nome + '.txt'), txt.strip())
    return jsonify({'ok': True, 'nome': nome, 'chars': len(txt)})

@app.route('/api/contexto/url', methods=['POST'])
def api_ctx_url():
    d = request.get_json(force=True) or {}
    emp, area, url = d.get('empresa', ''), d.get('area', ''), (d.get('url', '') or '').strip()
    if not re.match(r'^[\w-]+$', emp) or not re.match(r'^[\w-]+$', area):
        return jsonify({'erro': 'escolha empresa e área'}), 400
    if not url.startswith('http'): return jsonify({'erro': 'URL inválida'}), 400
    try:
        r = requests.get(url, timeout=30, headers={'User-Agent': 'Mozilla/5.0'})
        txt = extrair_texto('pagina.html', r.content)
    except Exception as e: return jsonify({'erro': 'falha ao buscar: %s' % e}), 400
    nome = 'site_' + re.sub(r'[^\w.-]+', '_', re.sub(r'^https?://', '', url))[:50]
    gravar(p(CTX, emp, area, nome + '.txt'), (url + '\n' + txt).strip())
    return jsonify({'ok': True, 'nome': nome, 'chars': len(txt)})

@app.route('/api/contexto/doc/<emp>/<area>/<nome>', methods=['DELETE'])
def api_ctx_del(emp, area, nome):
    if not all(re.match(r'^[\w.-]+$', x) for x in (emp, area, nome)): return ('', 404)
    f = p(CTX, emp, area, nome + '.txt')
    if os.path.exists(f): os.remove(f)
    return jsonify({'ok': True})

def registrar_conversa(pergunta, resposta, emp_nome='', emp_slug='', area=''):
    conv = json.loads(ler(p('conversas.json')) or '[]')
    item = {'id': max([c['id'] for c in conv], default=0) + 1,
            'data': datetime.datetime.now().isoformat(timespec='seconds'),
            'empresa': emp_nome, 'empresa_slug': emp_slug, 'area': area,
            'pergunta': pergunta, 'resposta': resposta}
    conv.append(item)
    gravar(p('conversas.json'), json.dumps(conv, ensure_ascii=False))
    return item

@app.route('/api/chat', methods=['POST'])
def api_chat():
    d = request.get_json(force=True)
    emp_slug, area = d.get('empresa') or None, d.get('area') or None
    try:
        resp = chat(d.get('pergunta', ''), d.get('historico', []), emp_slug, area)
        e = _empresa_por_slug(emp_slug)
        item = registrar_conversa(d.get('pergunta', ''), resp, e['nome'] if e else '', e['slug'] if e else '', area or '')
        return jsonify({'resposta': resp, 'id': item['id'], 'data': item['data'], 'empresa': item['empresa'], 'area': area or ''})
    except Exception as e:
        return jsonify({'erro': str(e)}), 500

@app.route('/api/empresas-areas')
def api_empresas_areas():
    garantir_empresa_padrao()
    out = []
    for e in empresas():
        arv = ctx_arvore(e['slug'])
        out.append({'slug': e['slug'], 'nome': e['nome'], 'ativa': e.get('ativa'),
                    'areas': [{'slug': a['slug'], 'nome': a['nome']} for a in arv['areas']]})
    return jsonify(out)

@app.route('/api/conversas')
def api_conversas():
    q = request.args.get('q', '').lower().strip()
    emp = request.args.get('empresa', '').strip()
    conv = sorted(json.loads(ler(p('conversas.json')) or '[]'), key=lambda c: c['id'], reverse=True)
    if emp:
        conv = [c for c in conv if c.get('empresa_slug', '') == emp]
    if q:
        conv = [c for c in conv if q in (c['pergunta'] + ' ' + c['resposta'] + ' ' + c.get('empresa', '')).lower()]
    return jsonify([{'id': c['id'], 'data': c['data'], 'empresa': c.get('empresa', ''), 'area': c.get('area', ''),
                     'pergunta': c['pergunta'], 'preview': c['resposta'][:180]} for c in conv])

@app.route('/api/conversa/<int:cid>')
def api_conversa(cid):
    conv = json.loads(ler(p('conversas.json')) or '[]')
    c = next((x for x in conv if x['id'] == cid), None)
    return (jsonify(c), 200) if c else (jsonify({}), 404)

@app.route('/api/conversa/<int:cid>', methods=['DELETE'])
def api_conversa_del(cid):
    conv = [x for x in json.loads(ler(p('conversas.json')) or '[]') if x['id'] != cid]
    gravar(p('conversas.json'), json.dumps(conv, ensure_ascii=False))
    return jsonify({'ok': True})

@app.route('/api/conversas', methods=['DELETE'])
def api_conversas_limpar():
    emp = request.args.get('empresa', '').strip()
    conv = json.loads(ler(p('conversas.json')) or '[]')
    if emp:
        conv = [c for c in conv if c.get('empresa_slug', '') != emp]
    else:
        conv = []
    gravar(p('conversas.json'), json.dumps(conv, ensure_ascii=False))
    return jsonify({'ok': True})

@app.route('/')
def painel():
    return Response(ler(os.path.join(BASE, 'painel.html')), mimetype='text/html')

# ---------- CRON ----------
try:
    from apscheduler.schedulers.background import BackgroundScheduler
    def rotina_cron():
        # padrão: só coleta os vídeos novos (você decide o que processar).
        # CRON_AUTO=1 no Railway → processa a fila inteira automaticamente.
        if os.environ.get('CRON_AUTO') == '1': processar()
        else:
            try: coletar()
            except Exception as e: print('cron coletar:', e, flush=True)
    sched = BackgroundScheduler(timezone='America/Sao_Paulo')
    sched.add_job(rotina_cron, 'cron', hour=CRON_HORA, minute=0)
    sched.start()
except Exception as e:
    print('APScheduler não iniciado:', e)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 8080)), debug=False)
